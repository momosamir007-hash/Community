import streamlit as st
from docx import Document
from cerebras.cloud.sdk import Cerebras
import pandas as pd
import json
import io
import os

# ---------------------------------------------------------
# 1. إعداد الصفحة وتصميمها
# ---------------------------------------------------------
st.set_page_config(
    page_title="مستخرج البيانات التربوية (Cerebras)",
    page_icon="🚀",
    layout="wide"
)

st.markdown("""
<style>
    .main { direction: rtl; text-align: right; }
    .stMarkdown, .stButton, .stDownloadButton, .stFileUploader, h1, h2, h3, p, div, label, input, .stSuccess { 
        text-align: right; 
        direction: rtl; 
    }
    .stDataFrame { direction: ltr; } 
    [data-testid="stSidebar"] { text-align: right; direction: rtl; }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# 2. إدارة المفاتيح (Secrets Management)
# ---------------------------------------------------------
with st.sidebar:
    st.header("⚙️ الإعدادات")
    
    # التحقق من وجود المفتاح في الأسرار
    if "CEREBRAS_API_KEY" in st.secrets:
        api_key = st.secrets["CEREBRAS_API_KEY"]
        st.success("✅ تم تحميل المفتاح من الأسرار (Secrets)")
    else:
        # إذا لم يكن في الأسرار، اطلبه من المستخدم
        api_key = st.text_input("Cerebras API Key", type="password")
        st.warning("لم يتم العثور على المفتاح في secrets.toml")

    # اختيار النموذج
    model_choice = st.selectbox(
        "اختر النموذج",
        ["llama3.1-70b", "llama-3.3-70b"],
        index=0
    )

# ---------------------------------------------------------
# 3. دوال المعالجة
# ---------------------------------------------------------
def extract_text_from_docx(file):
    """قراءة النصوص والجداول من ملف Word"""
    doc = Document(file)
    full_text = []
    
    # قراءة الفقرات
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # قراءة الجداول (دمج الصفوف بفاصل |)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                clean_text = cell.text.strip().replace("\n", " ")
                if clean_text:
                    row_text.append(clean_text)
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return "\n".join(full_text)

def analyze_with_cerebras(text, key, model_id):
    """تحليل النص باستخدام Cerebras API"""
    
    # تهيئة العميل
    client = Cerebras(api_key=key)
    
    system_prompt = """
    أنت مساعد تربوي خبير في تحليل المذكرات التربوية الجزائرية.
    استخرج البيانات التالية لكل نشاط/درس تجده في النص:
    1. "النشاط": (مثل: تعبير شفوي، رياضيات، تربية إسلامية...)
    2. "الموضوع": (عنوان الدرس)
    3. "الكفاءة_القاعدية": (نص الكفاءة)
    4. "مؤشر_الكفاءة": (المؤشر التربوي)

    القواعد:
    - المخرج يجب أن يكون JSON Valid (List of Objects).
    - إذا كانت المعلومة مفقودة اكتب "غير مذكور".
    - لا تضف أي شرح، فقط الـ JSON.
    """

    user_prompt = f"استخرج البيانات من هذا النص:\n{text[:25000]}"

    try:
        completion = client.chat.completions.create(
            model=model_id,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1,
            max_tokens=4000,
            response_format={"type": "json_object"}
        )
        
        response_content = completion.choices[0].message.content
        return json.loads(response_content)
        
    except Exception as e:
        return {"error": str(e)}

# ---------------------------------------------------------
# 4. الواجهة الرئيسية
# ---------------------------------------------------------
st.title("🚀 استخراج المذكرات (Cerebras AI)")
st.markdown("استخراج: **النشاط، الموضوع، الكفاءة، المؤشر** من ملفات Word.")

uploaded_file = st.file_uploader("📂 اختر ملف المذكرات (.docx)", type=["docx"])

if uploaded_file:
    if not api_key:
        st.error("⛔ يرجى توفير مفتاح API للمتابعة.")
    else:
        if st.button("⚡ ابدأ التحليل"):
            with st.spinner('جاري التحليل بسرعة فائقة...'):
                try:
                    # 1. استخراج النص
                    raw_text = extract_text_from_docx(uploaded_file)
                    
                    # 2. تحليل الذكاء الاصطناعي
                    result = analyze_with_cerebras(raw_text, api_key, model_choice)
                    
                    # 3. معالجة النتائج
                    data_list = []
                    if isinstance(result, list):
                        data_list = result
                    elif isinstance(result, dict):
                        # محاولة استخراج القائمة من داخل كائن JSON
                        for val in result.values():
                            if isinstance(val, list):
                                data_list = val
                                break
                        if not data_list: data_list = [result]

                    if data_list and "error" not in result:
                        st.success(f"تم استخراج {len(data_list)} نشاط!")
                        
                        # العرض في جدول
                        df = pd.DataFrame(data_list)
                        
                        # ترتيب الأعمدة المفضل
                        preferred_cols = ["النشاط", "الموضوع", "الكفاءة_القاعدية", "مؤشر_الكفاءة"]
                        final_cols = [c for c in preferred_cols if c in df.columns] + [c for c in df.columns if c not in preferred_cols]
                        df = df[final_cols]
                        
                        st.dataframe(df, use_container_width=True)
                        
                        # التحميل
                        col1, col2 = st.columns(2)
                        
                        # Excel
                        buffer = io.BytesIO()
                        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                            df.to_excel(writer, index=False, sheet_name='Data')
                        col1.download_button("📥 تحميل Excel", buffer.getvalue(), "lessons.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                        
                        # JSON
                        col2.download_button("📥 تحميل JSON", json.dumps(data_list, ensure_ascii=False, indent=4), "lessons.json", "application/json")
                        
                    elif "error" in result:
                        st.error(f"خطأ من المصدر: {result['error']}")
                    else:
                        st.warning("لم يتم العثور على بيانات.")
                        
                except Exception as e:
                    st.error(f"حدث خطأ: {e}")

import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION
from cerebras.cloud.sdk import Cerebras
import pandas as pd
import json
import io
import time

# ---------------------------------------------------------
# 1. إعداد الصفحة وتصميمها (CSS عربي وعصري)
# ---------------------------------------------------------
st.set_page_config(
    page_title="المساعد التربوي الشامل",
    page_icon="📚",
    layout="wide"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@400;500;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Tajawal', sans-serif;
        direction: rtl;
        text-align: right;
    }
    
    /* تنسيق العناوين */
    h1, h2, h3 { color: #154360; font-weight: 800; }
    
    /* تنسيق التبويبات */
    .stTabs [data-baseweb="tab-list"] { gap: 10px; }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #F0F3F4;
        border-radius: 5px 5px 0 0;
        color: #154360;
        font-weight: bold;
    }
    .stTabs [aria-selected="true"] {
        background-color: #154360;
        color: white;
    }

    /* الأزرار */
    .stButton>button {
        background-color: #154360;
        color: white;
        border-radius: 8px;
        font-weight: bold;
        transition: 0.3s;
    }
    .stButton>button:hover { background-color: #1A5276; }
    
    /* الجداول */
    [data-testid="stDataFrame"] { direction: rtl; }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# 2. بيانات التوقيت الأسبوعي (الثابتة)
# ---------------------------------------------------------
WEEKLY_SCHEDULE = {
    "الأحد": [
        {"time": "08:00 - 09:45", "activity": "تعبير شفوي"},
        {"time": "08:00 - 09:45", "activity": "مبادئ القراءة"},
        {"time": "08:00 - 09:45", "activity": "رياضيات"},
        {"time": "10:00 - 11:15", "activity": "ت علمية"},
        {"time": "10:00 - 11:15", "activity": "ت إسلامية"},
        {"time": "13:00 - 15:00", "activity": "مسرح وعرائس"},
        {"time": "13:00 - 15:00", "activity": "رسم وأشغال"},
        {"time": "13:00 - 15:00", "activity": "ت بدنية"}
    ],
    "الاثنين": [
        {"time": "08:00 - 09:45", "activity": "رياضيات"},
        {"time": "08:00 - 09:45", "activity": "تعبير شفوي"},
        {"time": "08:00 - 09:45", "activity": "تخطيط"},
        {"time": "10:00 - 11:15", "activity": "ت علمية"},
        {"time": "10:00 - 11:15", "activity": "ت مدنية"},
        {"time": "13:00 - 15:00", "activity": "مسرح وعرائس"},
        {"time": "13:00 - 15:00", "activity": "رسم وأشغال"},
        {"time": "13:00 - 15:00", "activity": "ت بدنية"}
    ],
    "الثلاثاء": [
        {"time": "08:00 - 09:45", "activity": "تعبير شفوي"},
        {"time": "08:00 - 09:45", "activity": "مبادئ القراءة"},
        {"time": "08:00 - 09:45", "activity": "رياضيات"},
        {"time": "10:00 - 11:15", "activity": "ت إسلامية"},
        {"time": "10:00 - 11:15", "activity": "ت بدنية"}
    ],
    "الأربعاء": [
        {"time": "08:00 - 09:45", "activity": "رياضيات"},
        {"time": "08:00 - 09:45", "activity": "مبادئ القراءة"},
        {"time": "08:00 - 09:45", "activity": "تخطيط"},
        {"time": "10:00 - 11:15", "activity": "ت علمية"},
        {"time": "10:00 - 11:15", "activity": "ت مدنية"},
        {"time": "13:00 - 15:00", "activity": "ت إيقاعية"},
        {"time": "13:00 - 15:00", "activity": "موسيقى وإنشاد"},
        {"time": "13:00 - 15:00", "activity": "ت بدنية"}
    ],
    "الخميس": [
        {"time": "08:00 - 09:45", "activity": "مبادئ القراءة"},
        {"time": "08:00 - 09:45", "activity": "رياضيات"},
        {"time": "08:00 - 09:45", "activity": "ت علمية"},
        {"time": "10:00 - 11:15", "activity": "ت إيقاعية"},
        {"time": "10:00 - 11:15", "activity": "موسيقى وإنشاد"}
    ]
}

# ---------------------------------------------------------
# 3. الدوال المساعدة (استخراج، تحليل، Word)
# ---------------------------------------------------------

def get_domain(activity):
    """تحديد المجال تلقائياً بناءً على النشاط"""
    act = activity.strip()
    if any(x in act for x in ["تعبير", "قراءة", "تخطيط", "لغة"]): return "اللغوي"
    elif "رياضيات" in act: return "الرياضي"
    elif any(x in act for x in ["علمية", "تكنولوجيا"]): return "العلمي"
    elif any(x in act for x in ["إسلامية", "مدنية"]): return "الاجتماعي"
    elif any(x in act for x in ["مسرح", "رسم", "موسيقى", "إنشاد", "تشكيلية"]): return "الفني"
    elif any(x in act for x in ["بدنية", "إيقاعية", "رياضة"]): return "البدني والإيقاعي"
    return ""

def extract_text_from_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip(): full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [c.text.strip().replace("\n", " ") for c in row.cells if c.text.strip()]
            if row_text: full_text.append(" | ".join(row_text))
    return "\n".join(full_text)

def analyze_with_cerebras(text, key, model_id, mode="journal"):
    """
    mode='journal': استخراج بسيط للمذكرة اليومية (نشاط، موضوع، كفاءة).
    mode='structure': استخراج عميق للعناوين والمقاطع (المجال، الوحدة، إلخ).
    """
    client = Cerebras(api_key=key)
    
    if mode == "journal":
        prompt = """
        استخرج بيانات الدروس لملء جدول يومي.
        المطلوب JSON List للكائنات:
        {"النشاط": "...", "الموضوع": "...", "الكفاءة": "...", "المؤشر": "..."}
        حاول مطابقة أسماء الأنشطة مع (رياضيات، تعبير شفوي، ت علمية، إلخ).
        """
    else: # mode == structure
        prompt = """
        قم بتحليل النص لاستخراج الهيكلة الكاملة للدروس.
        المطلوب JSON List للكائنات:
        1. "المجال_أو_المقطع": (العنوان الكبير، مثل: المجال اللغوي، الوحدة 3، الحياة المدرسية).
        2. "النشاط": نوع الحصة.
        3. "الموضوع": عنوان الدرس.
        4. "الكفاءة_الختامية": الكفاءة.
        5. "المؤشر": الهدف التعلمي.
        """

    try:
        completion = client.chat.completions.create(
            model=model_id,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": text[:28000]}
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        return json.loads(completion.choices[0].message.content)
    except Exception as e:
        return {"error": str(e)}

def create_daily_journal_doc(day_name, extracted_lessons):
    """إنشاء ملف Word للمذكرة اليومية"""
    doc = Document()
    
    # إعداد الصفحة Landscape
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.orientation = 1 
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # العنوان
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'المذكرة اليومية - يوم: {day_name}')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(21, 67, 96) # Dark Blue

    # الجدول
    headers = ["التوقيت", "النشاط", "المجال", "الموضوع (المحتوى)", "الكفاءة", "المؤشر", "ملاحظات"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.direction = WD_TABLE_DIRECTION.RTL
    table.autofit = False 
    
    # تنسيق الرأس
    widths = [0.8, 1.0, 0.9, 1.5, 1.2, 1.2, 0.8]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.width = Inches(widths[i])
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(11)
        
    # تعبئة البيانات
    day_schedule = WEEKLY_SCHEDULE.get(day_name, [])
    
    # معالجة البيانات المستخرجة
    lessons_list = []
    if isinstance(extracted_lessons, dict):
        for val in extracted_lessons.values():
            if isinstance(val, list): lessons_list = val; break
        if not lessons_list: lessons_list = [extracted_lessons]
    else: lessons_list = extracted_lessons

    for slot in day_schedule:
        row = table.add_row()
        cells = row.cells
        
        # 1. التوقيت والنشاط (ثابت)
        cells[0].text = slot['time']
        cells[1].text = slot['activity']
        
        # 2. المجال (تلقائي)
        cells[2].text = get_domain(slot['activity'])
        
        # 3. البحث عن الدرس (AI Data)
        found = None
        clean_slot = slot['activity'].replace("ت ", "").replace("مبادئ ", "").strip()
        for lesson in lessons_list:
            lesson_act = str(lesson.get('النشاط', '')).replace("ت ", "").replace("مبادئ ", "").strip()
            if clean_slot in lesson_act or lesson_act in clean_slot:
                found = lesson
                break
        
        if found:
            cells[3].text = str(found.get('الموضوع', ''))
            cells[4].text = str(found.get('الكفاءة', '') or found.get('الكفاءة_الختامية', ''))
            cells[5].text = str(found.get('المؤشر', ''))
        
        # تنسيق الخلايا
        for i, cell in enumerate(cells):
            cell.width = Inches(widths[i])
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if p.runs: p.runs[0].font.name = "Arial"

    return doc

# ---------------------------------------------------------
# 4. الواجهة الرئيسية
# ---------------------------------------------------------
with st.sidebar:
    st.title("⚙️ الإعدادات")
    api_key = st.secrets.get("CEREBRAS_API_KEY") or st.text_input("Cerebras API Key", type="password")
    model_choice = st.selectbox("النموذج", ["llama-3.3-70b", "llama3.1-8b"])
    st.info("💡 استخدم Llama 3.3 لأفضل نتائج.")

st.title("📚 المنصة التربوية الذكية")

# --- نظام التبويبات ---
tab1, tab2 = st.tabs(["📝 المذكرة اليومية (ملء التوقيت)", "📊 استخراج العناوين والجدول"])

# ==========================================
# Tab 1: المذكرة اليومية
# ==========================================
with tab1:
    st.header("إنشاء مذكرة يومية حسب التوقيت الأسبوعي")
    col1, col2 = st.columns([2, 1])
    
    with col1:
        file_tab1 = st.file_uploader("📂 ملف المذكرات (.docx)", type=["docx"], key="f1")
    with col2:
        day_selected = st.selectbox("📅 اختر اليوم:", list(WEEKLY_SCHEDULE.keys()))

    if file_tab1 and st.button("🚀 إنشاء المذكرة", key="btn1"):
        if not api_key: st.error("المفتاح مفقود!"); st.stop()
        
        with st.spinner(f'جاري تحليل الدروس ومطابقتها مع توقيت يوم {day_selected}...'):
            text = extract_text_from_docx(file_tab1)
            data = analyze_with_cerebras(text, api_key, model_choice, mode="journal")
            
            if "error" not in data:
                doc = create_daily_journal_doc(day_selected, data)
                bio = io.BytesIO()
                doc.save(bio)
                
                st.success("✅ تم إنشاء الملف بنجاح! تم تحديد المجالات تلقائياً.")
                st.download_button(
                    label=f"📥 تحميل مذكرة {day_selected} (Word)",
                    data=bio.getvalue(),
                    file_name=f"Journal_{day_selected}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error(f"خطأ: {data['error']}")

# ==========================================
# Tab 2: استخراج العناوين (المجال/المقطع)
# ==========================================
with tab2:
    st.header("استخراج البيانات التربوية (Excel/JSON)")
    st.markdown("استخراج: **المجال/المقطع، النشاط، الموضوع، الكفاءة** في جدول منظم.")
    
    file_tab2 = st.file_uploader("📂 ملف المذكرات (.docx)", type=["docx"], key="f2")
    
    if file_tab2 and st.button("🔍 تحليل واستخراج البيانات", key="btn2"):
        if not api_key: st.error("المفتاح مفقود!"); st.stop()
        
        with st.spinner('جاري استخراج العناوين والهيكلة...'):
            text = extract_text_from_docx(file_tab2)
            # وضع 'structure' هنا لتفعيل استخراج العناوين والمجالات
            result = analyze_with_cerebras(text, api_key, model_choice, mode="structure")
            
            # معالجة JSON
            final_data = []
            if isinstance(result, dict):
                for val in result.values():
                    if isinstance(val, list): final_data = val; break
                if not final_data and "error" not in result: final_data = [result]
            elif isinstance(result, list): final_data = result
            
            if final_data:
                df = pd.DataFrame(final_data)
                
                # ترتيب الأعمدة
                cols = ["المجال_أو_المقطع", "النشاط", "الموضوع", "الكفاءة_الختامية", "المؤشر"]
                df = df[[c for c in cols if c in df.columns] + [c for c in df.columns if c not in cols]]
                
                st.success(f"تم استخراج {len(df)} درساً.")
                
                # جدول قابل للتعديل
                st.subheader("📝 مراجعة وتعديل البيانات")
                edited_df = st.data_editor(df, use_container_width=True, num_rows="dynamic")
                
                # تحميل
                buffer = io.BytesIO()
                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                    edited_df.to_excel(writer, index=False, sheet_name='Data')
                
                c1, c2 = st.columns(2)
                c1.download_button("📥 تحميل Excel", buffer.getvalue(), "lessons_structured.xlsx")
                c2.download_button("📥 تحميل JSON", json.dumps(final_data, ensure_ascii=False), "data.json")
                
            else:
                st.error("لم يتم العثور على بيانات مهيكلة أو حدث خطأ.")
